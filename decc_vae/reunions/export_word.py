# Compatibilité : l'export Word détaillé reste ici ; le rapport utilise export_utils.
from io import BytesIO

from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x0B, 0x25, 0x45)
RED = RGBColor(0xC0, 0x39, 0x2B)


def _set_run_font(run, size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def export_reunion_docx(reunion):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = header.add_run("DECC / VAE — Ministère de l'Éducation Nationale (Niger)")
        _set_run_font(r, size=10, bold=True, color=NAVY)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer.add_run(
            f"Document généré le {timezone.localtime(timezone.now()).strftime('%d/%m/%Y %H:%M')}"
        )
        _set_run_font(fr, size=9, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = p.add_run("Compte rendu de réunion")
    _set_run_font(t, 16, True, NAVY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(p2.add_run(str(reunion)), 13, True)

    rows = [
        ("Date", reunion.date.strftime("%d/%m/%Y")),
        (
            "Horaire",
            f"{reunion.heure_debut:%H:%M}"
            + (f" – {reunion.heure_fin:%H:%M}" if reunion.heure_fin else ""),
        ),
        ("Lieu", reunion.get_lieu_label()),
        ("Type", reunion.get_type_label()),
        ("Président", reunion.president),
        ("Rapporteur", reunion.rapporteur),
        ("Nombre de participants", str(reunion.nombre_participants)),
        (
            "Prochaine réunion",
            reunion.prochaine_reunion.strftime("%d/%m/%Y") if reunion.prochaine_reunion else "—",
        ),
    ]
    info = doc.add_table(rows=len(rows), cols=2)
    info.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        info.rows[i].cells[0].text = ""
        info.rows[i].cells[1].text = ""
        _set_run_font(info.rows[i].cells[0].paragraphs[0].add_run(label), 12, True)
        _set_run_font(info.rows[i].cells[1].paragraphs[0].add_run(value), 12, False)

    h = doc.add_paragraph()
    _set_run_font(h.add_run("Participants"), 13, True, NAVY)
    for label, text in (
        ("Présents", reunion.participants_presents),
        ("Excusés", reunion.participants_excuses),
        ("Absents", reunion.participants_absents),
    ):
        para = doc.add_paragraph()
        _set_run_font(para.add_run(f"{label} : "), 12, True)
        _set_run_font(para.add_run(text or "—"), 12, False)

    h = doc.add_paragraph()
    _set_run_font(h.add_run("Tableau des points"), 13, True, NAVY)
    headers = [
        "N°",
        "Rubrique",
        "Volet",
        "Sujet",
        "Décision",
        "Action",
        "Responsable",
        "Délai",
        "Urgence",
        "Statut",
    ]
    points = list(reunion.points.all())
    table = doc.add_table(rows=1 + len(points), cols=len(headers))
    table.style = "Table Grid"
    for i, htxt in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _set_run_font(cell.paragraphs[0].add_run(htxt), 10, True, RGBColor(255, 255, 255))
        cell._tc.get_or_add_tcPr().append(
            parse_xml(
                r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="0B2545"/>'
            )
        )
    for idx, point in enumerate(points, start=1):
        values = [
            str(point.numero),
            point.get_rubrique_display(),
            point.get_volet_label(),
            point.sujet or "",
            point.decision or "",
            point.action or "",
            point.responsable or "",
            point.delai.strftime("%d/%m/%Y") if point.delai else "—",
            str(point.urgence),
            point.get_statut_display(),
        ]
        for j, val in enumerate(values):
            cell = table.rows[idx].cells[j]
            cell.text = ""
            _set_run_font(
                cell.paragraphs[0].add_run(val),
                9,
                bold=(point.urgence == 5 and j == 8),
                color=RED if point.urgence == 5 else None,
            )
            if point.urgence == 5:
                cell._tc.get_or_add_tcPr().append(
                    parse_xml(
                        r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="FDECEA"/>'
                    )
                )

    h = doc.add_paragraph()
    _set_run_font(h.add_run("Synthèse"), 13, True, NAVY)
    n = len(points)
    n5 = sum(1 for x in points if x.urgence == 5)
    n_fait = sum(1 for x in points if x.statut == "fait")
    synth = doc.add_paragraph()
    _set_run_font(
        synth.add_run(
            (
                f"Cette réunion a examiné {n} point(s), dont {n5} critique(s). "
                f"{n_fait} point(s) terminé(s). {reunion.objet_prochaine or ''} {reunion.observations or ''}"
            ).strip()
            or "—"
        ),
        12,
    )
    sig = doc.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    sig.rows[0].cells[0].text = "Président"
    sig.rows[0].cells[1].text = "Rapporteur"
    sig.rows[1].cells[0].text = f"Nom : {reunion.president}\n\nSignature :\n\n"
    sig.rows[1].cells[1].text = f"Nom : {reunion.rapporteur}\n\nSignature :\n\n"
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
