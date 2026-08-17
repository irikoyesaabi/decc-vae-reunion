"""Exports PDF (xhtml2pdf), Word, Excel et import Excel."""
from io import BytesIO
from datetime import datetime

from django.template.loader import render_to_string
from django.utils import timezone
from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from xhtml2pdf import pisa

from .models import Point, Reunion

NAVY_FILL = PatternFill("solid", fgColor="0B2545")
NAVY_FONT = Font(color="FFFFFF", bold=True, name="Calibri")
RED_FILL = PatternFill("solid", fgColor="C0392B")
RED_FONT = Font(color="FFFFFF", bold=True)
THIN = Border(
    left=Side(style="thin", color="9AA8B8"),
    right=Side(style="thin", color="9AA8B8"),
    top=Side(style="thin", color="9AA8B8"),
    bottom=Side(style="thin", color="9AA8B8"),
)

VOLET_IMPORT = {label.lower(): code for code, label in Point.VOLET_CHOICES}
VOLET_IMPORT.update(
    {
        "si": Point.VOLET_DONNEES,
        "gestion des données": Point.VOLET_DONNEES,
        "gestion des donnees": Point.VOLET_DONNEES,
        "données": Point.VOLET_DONNEES,
    }
)
RUBRIQUE_IMPORT = {label.lower(): code for code, label in Point.RUBRIQUE_CHOICES}
STATUT_IMPORT = {label.lower(): code for code, label in Point.STATUT_CHOICES}


def _html_to_pdf(html):
    dest = BytesIO()
    result = pisa.CreatePDF(html, dest=dest, encoding="utf-8")
    if result.err:
        raise ValueError("Échec de la génération PDF.")
    return dest.getvalue()


def export_reunion_pdf(reunion):
    html = render_to_string(
        "reunions/synthese_pdf.html",
        {
            "reunion": reunion,
            "points": reunion.points.all(),
            "generated": timezone.localtime(timezone.now()),
            "mode": "reunion",
        },
    )
    return _html_to_pdf(html)


def export_rapport_pdf(reunions):
    reunions = list(reunions)
    points = []
    for r in reunions:
        points.extend(list(r.points.all()))
    html = render_to_string(
        "reunions/synthese_pdf.html",
        {
            "reunions": reunions,
            "points": points,
            "generated": timezone.localtime(timezone.now()),
            "mode": "rapport",
        },
    )
    return _html_to_pdf(html)


def export_reunion_docx(reunion):
    from .export_word import export_reunion_docx as _docx

    return _docx(reunion)


def export_rapport_word(reunions):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DECC / VAE — Rapport de réunions")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    doc.add_paragraph(f"Généré le {timezone.localtime(timezone.now()).strftime('%d/%m/%Y %H:%M')}")
    reunions = list(reunions)
    points = []
    for r in reunions:
        points.extend(list(r.points.all()))
    doc.add_paragraph(f"Nombre de réunions : {len(reunions)}")
    doc.add_paragraph(f"Nombre de points : {len(points)}")
    doc.add_paragraph(f"Points critiques : {sum(1 for x in points if x.urgence == 5)}")
    for reunion in reunions:
        doc.add_paragraph()
        h = doc.add_paragraph()
        r = h.add_run(str(reunion))
        r.bold = True
        for point in reunion.points.all():
            doc.add_paragraph(
                f"{point.numero}. [{point.get_volet_label()}] {point.sujet} "
                f"(urgence {point.urgence}, {point.get_statut_display()})"
            )
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _style_header(ws, ncol):
    for col in range(1, ncol + 1):
        cell = ws.cell(1, col)
        cell.fill = NAVY_FILL
        cell.font = NAVY_FONT
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = THIN
    ws.auto_filter.ref = f"A1:{get_column_letter(ncol)}{max(ws.max_row, 1)}"
    ws.freeze_panes = "A2"


def _autosize(ws):
    for col in ws.columns:
        letter = get_column_letter(col[0].column)
        length = 10
        for cell in col:
            val = "" if cell.value is None else str(cell.value)
            length = max(length, min(len(val) + 2, 45))
        ws.column_dimensions[letter].width = length


def _write_points(ws, points, with_reunion=False):
    headers = []
    if with_reunion:
        headers.append("Réunion")
    headers += [
        "N°",
        "Rubrique",
        "Volet",
        "Sujet",
        "Décision",
        "Action",
        "Responsable",
        "Délai",
        "Urgence",
        "Statut",
    ]
    ws.append(headers)
    for point in points:
        row = []
        if with_reunion:
            row.append(f"{point.reunion.date:%d/%m/%Y} — {point.reunion.get_type_label()}")
        row += [
            point.numero,
            point.get_rubrique_display(),
            point.get_volet_label(),
            point.sujet,
            point.decision,
            point.action,
            point.responsable,
            point.delai.strftime("%d/%m/%Y") if point.delai else "",
            point.urgence,
            point.get_statut_display(),
        ]
        ws.append(row)
    ncol = len(headers)
    _style_header(ws, ncol)
    urg_col = get_column_letter(ncol - 1)
    if ws.max_row >= 2:
        ws.conditional_formatting.add(
            f"{urg_col}2:{urg_col}{ws.max_row}",
            CellIsRule(operator="equal", formula=["5"], fill=RED_FILL, font=RED_FONT),
        )
    _autosize(ws)


def export_reunion_xlsx(reunion):
    wb = Workbook()
    ws = wb.active
    ws.title = "Points"
    _write_points(ws, reunion.points.all(), with_reunion=False)
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def export_all_reunions_xlsx(reunions_qs):
    reunions = list(reunions_qs.prefetch_related("points"))
    wb = Workbook()
    ws_r = wb.active
    ws_r.title = "Réunions"
    ws_r.append(["Date", "Type", "Président", "Nb participants", "Lieu", "Rapporteur"])
    for r in reunions:
        ws_r.append(
            [
                r.date.strftime("%d/%m/%Y"),
                r.get_type_label(),
                r.president,
                r.nombre_participants,
                r.get_lieu_label(),
                r.rapporteur,
            ]
        )
    _style_header(ws_r, 6)
    _autosize(ws_r)
    ws_p = wb.create_sheet("Points")
    all_points = []
    for r in reunions:
        all_points.extend(list(r.points.all()))
    _write_points(ws_p, all_points, with_reunion=True)
    ws_s = wb.create_sheet("Statistiques")
    points = all_points
    ws_s.append(["Indicateur", "Valeur"])
    ws_s.append(["Nombre de réunions", len(reunions)])
    ws_s.append(["Nombre de points", len(points)])
    ws_s.append(["Points critiques (urgence 5)", sum(1 for p in points if p.urgence == 5)])
    ws_s.append(["Points urgents (4 et 5)", sum(1 for p in points if p.urgence >= 4)])
    ws_s.append(["Points à faire", sum(1 for p in points if p.statut == Point.STATUT_A_FAIRE)])
    ws_s.append(["Points en cours", sum(1 for p in points if p.statut == Point.STATUT_EN_COURS)])
    ws_s.append(["Points faits", sum(1 for p in points if p.statut == Point.STATUT_FAIT)])
    ws_s.append([])
    ws_s.append(["Points par volet", "Nombre"])
    for code, label in Point.VOLET_CHOICES:
        ws_s.append([label, sum(1 for p in points if p.volet == code)])
    _style_header(ws_s, 2)
    _autosize(ws_s)
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def excel_template():
    wb = Workbook()
    ws = wb.active
    ws.title = "Points"
    headers = [
        "Rubrique",
        "Volet",
        "Sujet",
        "Décision",
        "Action",
        "Responsable",
        "Délai",
        "Urgence",
        "Statut",
    ]
    ws.append(headers)
    ws.append(
        [
            "ODJ",
            "Examens",
            "Exemple de sujet",
            "Décision",
            "Action",
            "Nom du responsable",
            datetime.today().strftime("%Y-%m-%d"),
            3,
            "À faire",
        ]
    )
    _style_header(ws, len(headers))
    _autosize(ws)
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _parse_date(value):
    if value is None or value == "":
        return None
    if hasattr(value, "date"):
        return value.date() if not isinstance(value, datetime) else value.date()
    if isinstance(value, datetime):
        return value.date()
    text = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            continue
    return None


def import_points_from_xlsx(reunion, uploaded_file):
    wb = load_workbook(uploaded_file, data_only=True)
    ws = wb.active
    created = 0
    errors = []
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for i, row in enumerate(rows, start=2):
        if not row or all(c is None or str(c).strip() == "" for c in row):
            continue
        rubrique_raw = str(row[0] or "").strip().lower()
        volet_raw = str(row[1] or "").strip().lower()
        sujet = str(row[2] or "").strip()
        if not sujet:
            errors.append(f"Ligne {i} : sujet manquant.")
            continue
        rubrique = RUBRIQUE_IMPORT.get(rubrique_raw, Point.RUBRIQUE_ODJ)
        volet = VOLET_IMPORT.get(volet_raw, Point.VOLET_AUTRE)
        volet_prec = ""
        if volet == Point.VOLET_AUTRE and volet_raw not in ("autre", ""):
            volet_prec = str(row[1] or "").strip()
        statut_raw = str(row[8] or "").strip().lower() if len(row) > 8 else ""
        try:
            urgence = int(row[7] or 3)
        except (TypeError, ValueError):
            urgence = 3
        urgence = min(5, max(1, urgence))
        Point.objects.create(
            reunion=reunion,
            numero=0,
            rubrique=rubrique,
            volet=volet,
            volet_autre_precision=volet_prec,
            sujet=sujet,
            decision=str(row[3] or "").strip(),
            action=str(row[4] or "").strip(),
            responsable=str(row[5] or "").strip(),
            delai=_parse_date(row[6] if len(row) > 6 else None),
            urgence=urgence,
            statut=STATUT_IMPORT.get(statut_raw, Point.STATUT_A_FAIRE),
        )
        created += 1
    return created, errors
